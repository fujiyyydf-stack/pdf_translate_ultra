#!/usr/bin/env python3
"""
Word 文档处理模块
功能：读取、解析 Word 文档，提取段落文本
"""

import re
from pathlib import Path
from typing import List, Dict, Optional

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，Word 文件处理功能将不可用")
    print("请运行: pip install python-docx")


# 支持的文件格式
SUPPORTED_FORMATS = ['.docx']


def is_supported_word_file(filepath: str) -> bool:
    """检查是否为支持的 Word 文件格式"""
    return Path(filepath).suffix.lower() in SUPPORTED_FORMATS


class WordProcessor:
    """Word 文档处理器"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")
    
    def extract_paragraphs(self, docx_path: str) -> List[Dict]:
        """
        从 Word 文档提取段落
        
        Args:
            docx_path: Word 文件路径
            
        Returns:
            段落列表 [{"index": 0, "text": "...", "style": "Normal", ...}]
        """
        doc = Document(docx_path)
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:  # 忽略空段落
                paragraphs.append({
                    "index": i,
                    "text": text,
                    "style": para.style.name if para.style else "Normal",
                    "is_heading": para.style and "Heading" in para.style.name if para.style else False,
                    "char_count": len(text)
                })
        
        return paragraphs
    
    def extract_with_formatting(self, docx_path: str) -> List[Dict]:
        """
        提取段落，保留更多格式信息（用于精确匹配）
        
        Args:
            docx_path: Word 文件路径
            
        Returns:
            带格式信息的段落列表
        """
        doc = Document(docx_path)
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # 检测格式特征
            is_bold = any(run.bold for run in para.runs if run.bold)
            is_italic = any(run.italic for run in para.runs if run.italic)
            
            # 检测是否为标题
            is_heading = False
            heading_level = 0
            if para.style:
                style_name = para.style.name
                is_heading = "Heading" in style_name or "标题" in style_name
                # 提取标题级别
                level_match = re.search(r'(\d+)', style_name)
                if level_match:
                    heading_level = int(level_match.group(1))
            
            paragraphs.append({
                "index": i,
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "is_heading": is_heading,
                "heading_level": heading_level,
                "is_bold": is_bold,
                "is_italic": is_italic,
                "char_count": len(text),
                "word_count": len(text.split())
            })
        
        return paragraphs
    
    def extract_by_sections(self, docx_path: str) -> List[Dict]:
        """
        按章节提取文本（适用于有明确章节结构的文档）
        
        Returns:
            章节列表 [{"title": "...", "level": 1, "paragraphs": [...]}]
        """
        paragraphs = self.extract_with_formatting(docx_path)
        
        sections = []
        current_section = {
            "title": "开篇",
            "level": 0,
            "paragraphs": []
        }
        
        for para in paragraphs:
            if para["is_heading"]:
                # 保存当前章节
                if current_section["paragraphs"]:
                    sections.append(current_section)
                
                # 开始新章节
                current_section = {
                    "title": para["text"],
                    "level": para["heading_level"],
                    "paragraphs": []
                }
            else:
                current_section["paragraphs"].append(para)
        
        # 保存最后一个章节
        if current_section["paragraphs"]:
            sections.append(current_section)
        
        return sections
    
    def merge_short_paragraphs(
        self, 
        paragraphs: List[Dict], 
        min_length: int = 50
    ) -> List[Dict]:
        """
        合并过短的段落（可能是被错误分割的）
        
        Args:
            paragraphs: 段落列表
            min_length: 最小长度阈值
            
        Returns:
            合并后的段落列表
        """
        if not paragraphs:
            return []
        
        merged = []
        buffer = None
        
        for para in paragraphs:
            # 标题不合并
            if para.get("is_heading"):
                if buffer:
                    merged.append(buffer)
                    buffer = None
                merged.append(para)
                continue
            
            if buffer is None:
                buffer = para.copy()
            elif len(buffer["text"]) < min_length:
                # 合并到 buffer
                buffer["text"] += "\n" + para["text"]
                buffer["char_count"] = len(buffer["text"])
            else:
                merged.append(buffer)
                buffer = para.copy()
        
        if buffer:
            merged.append(buffer)
        
        # 重新编号
        for i, para in enumerate(merged):
            para["merged_index"] = i
        
        return merged
    
    def get_document_stats(self, docx_path: str) -> Dict:
        """
        获取文档统计信息
        
        Returns:
            统计信息字典
        """
        paragraphs = self.extract_paragraphs(docx_path)
        
        total_chars = sum(p["char_count"] for p in paragraphs)
        total_words = sum(len(p["text"].split()) for p in paragraphs)
        heading_count = sum(1 for p in paragraphs if p.get("is_heading"))
        
        return {
            "total_paragraphs": len(paragraphs),
            "total_characters": total_chars,
            "total_words": total_words,
            "heading_count": heading_count,
            "avg_paragraph_length": total_chars // len(paragraphs) if paragraphs else 0
        }
    
    def extract_text_only(self, docx_path: str) -> str:
        """
        提取纯文本（用于简单场景）
        
        Returns:
            完整文本字符串
        """
        doc = Document(docx_path)
        texts = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)
        
        return "\n\n".join(texts)


def test_word_processor():
    """测试函数"""
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python word_processor.py <word_file.docx>")
        return
    
    filepath = sys.argv[1]
    
    if not Path(filepath).exists():
        print(f"文件不存在: {filepath}")
        return
    
    processor = WordProcessor()
    
    # 获取统计信息
    stats = processor.get_document_stats(filepath)
    print("\n📊 文档统计:")
    print(f"  段落数: {stats['total_paragraphs']}")
    print(f"  字符数: {stats['total_characters']}")
    print(f"  标题数: {stats['heading_count']}")
    print(f"  平均段落长度: {stats['avg_paragraph_length']}")
    
    # 提取段落
    paragraphs = processor.extract_with_formatting(filepath)
    print(f"\n📝 前 5 个段落:")
    for para in paragraphs[:5]:
        text_preview = para["text"][:50] + "..." if len(para["text"]) > 50 else para["text"]
        print(f"  [{para['index']}] {para['style']}: {text_preview}")


if __name__ == "__main__":
    test_word_processor()
